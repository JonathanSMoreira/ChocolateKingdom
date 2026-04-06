# -*- coding: utf-8 -*-
"""
Gera DOCFUNCIONALIDADES_PT-BR.docx e DOCFUNCIONALIDADES_EN.docx.
Executar na raiz do projeto:
  python docs/generate_funcionalidades_word.py
Requisito:
  pip install python-docx
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def set_doc_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_heading(text, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_p(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def build_pt() -> Document:
    doc = Document()
    set_doc_defaults(doc)
    add_title(doc, "Cacau Parque — Documentacao do projeto")

    add_p(
        doc,
        "Este documento foi escrito para dois publicos: no inicio, linguagem acessivel para pessoas leigas; "
        "no final, visao tecnica para engenharia, manutencao e evolucao do produto.",
    )

    add_h(doc, "Parte 1 — Explicacao simples (leigos)", 1)
    add_h(doc, "Objetivo do app", 2)
    add_p(
        doc,
        "O app Cacau Parque foi criado para apresentar a experiencia de um parque tematico no celular: "
        "explorar atracoes, visualizar mapa, criar conta, fazer login e gerir perfil.",
    )
    add_h(doc, "O que ele faz hoje", 2)
    add_p(doc, "Home: vitrine do parque com destaques de lojas, eventos e atracoes.")
    add_p(doc, "Ingressos: area visual, ainda em evolucao para compra real.")
    add_p(doc, "Mapa: visualizacao com zoom e pontos tocaveis com informacoes.")
    add_p(doc, "Perfil: login, cadastro, edicao de dados, foto, endereco e saida da conta.")
    add_h(doc, "Linguagens e tecnologias (em termos simples)", 2)
    add_p(doc, "Celular: React Native com Expo (app multiplataforma).")
    add_p(doc, "Servidor: Node.js com Express (responde as requisicoes).")
    add_p(doc, "Banco de dados: Microsoft SQL Server (onde os dados ficam guardados).")

    add_h(doc, "Parte 2 — Funcionalidades por pagina", 1)
    add_h(doc, "Home", 2)
    add_p(
        doc,
        "Tela inicial com identidade visual, secoes em carrossel e atalhos que encaminham para o mapa filtrado.",
    )
    add_h(doc, "Ingressos", 2)
    add_p(doc, "Tela de placeholder com identidade visual; compra ainda nao habilitada nesta versao.")
    add_h(doc, "Mapa", 2)
    add_p(
        doc,
        "Gestos de arrasto e zoom, hotspots com detalhes e fallback para dados locais quando API falha.",
    )
    add_h(doc, "Perfil", 2)
    add_p(
        doc,
        "Fluxos de autenticacao, dados pessoais e recursos extras para perfis marcados como funcionarios.",
    )

    add_h(doc, "Parte 3 — Arquitetura, engenharia e estado das areas", 1)
    add_h(doc, "Arquitetura geral", 2)
    add_p(
        doc,
        "Modelo cliente-servidor: aplicativo React Native consome API REST em Node/Express, "
        "que persiste dados no SQL Server.",
    )
    add_h(doc, "Estado atual por area", 2)
    add_p(
        doc,
        "Backend: mais proximo de arranjo profissional, com routes/, services/, utils/, schema/, workers/ e db/connection.js. "
        "Ainda existe concentracao relevante no server.js.",
    )
    add_p(
        doc,
        "App React Native: houve extracao parcial para src/ (config, types, services/http, utils, map/, screens/MapaTab.tsx), "
        "mas App.tsx continua grande e concentra muitas responsabilidades.",
    )
    add_p(
        doc,
        "Confiabilidade/manutencao: melhorou com health check, timeout, troubleshooting SQL e separacao de .env. "
        "Proximo passo natural: modularizar por ecras e hooks.",
    )
    add_h(doc, "Tipos de linguagem usados no projeto", 2)
    add_p(doc, "TypeScript e JavaScript no app e backend.")
    add_p(doc, "SQL para schema, migracoes e seeds.")
    add_p(doc, "PowerShell/BAT para automacoes de ambiente SQL e suporte operacional.")

    add_h(doc, "Parte 4 — Avaliacao honesta: esta nivel profissional?", 1)
    add_p(
        doc,
        "Em partes sim, no todo ainda nao no sentido de app totalmente dividido por paginas. "
        "E um projeto serio e viavel para portfolio e evolucao incremental.",
    )

    add_h(doc, "Visao directa (o que ainda pesa)", 2)
    add_p(
        doc,
        "1) Frontend: App.tsx monolitico (milhares de linhas). Impacto: manutencao e revisoes mais dificeis. "
        "Direcao: extrair screens, hooks e componentes compartilhados.",
    )
    add_p(
        doc,
        "2) Backend: server.js ainda grande. Impacto: crescimento no mesmo ficheiro. "
        "Direcao: separar rotas por dominio (auth, clientes, funcionarios, etc.).",
    )
    add_p(
        doc,
        "3) Testes automatizados pouco visiveis. Impacto: refactors com maior risco. "
        "Direcao: testes de API para endpoints criticos e smoke tests no cliente.",
    )
    add_p(
        doc,
        "4) Configuracao/ambientes: multiplos modos SQL e variaveis. Impacto: onboarding sujeito a erro. "
        "Direcao: guia unico de setup com variaveis obrigatorias.",
    )
    add_p(
        doc,
        "5) SQL no repositorio: muitos scripts historicos. Impacto: risco de executar script errado. "
        "Direcao: padronizar bootstrap e cabecalhos claros.",
    )
    add_p(
        doc,
        "6) Observabilidade: logs funcionais, mas pouco estruturados. Direcao: padrao de logs e erros centralizados.",
    )
    add_p(
        doc,
        "7) Seguranca operacional: manter .env fora do Git, revisar secrets e reforcar HTTPS em producao.",
    )

    add_h(doc, "Resumo final (PT-BR)", 2)
    add_p(doc, "Arquitetura geral: backend mais organizado que o app.")
    add_p(doc, "Manutencao: principal divida tecnica esta no App.tsx e no server.js.")
    add_p(doc, "Profissionalismo: solido para portfolio e producao incremental, com espaco claro para modularizacao.")

    return doc


def build_en() -> Document:
    doc = Document()
    set_doc_defaults(doc)
    add_title(doc, "Cacau Parque — Project documentation")

    add_p(
        doc,
        "This document is split for two audiences: plain-language explanation first, then a technical section "
        "covering architecture, engineering state, and codebase maturity.",
    )

    add_h(doc, "Part 1 — Plain-language overview", 1)
    add_h(doc, "App objective", 2)
    add_p(
        doc,
        "Cacau Parque is a mobile app designed to present a theme-park experience: explore attractions, "
        "view a map, create an account, sign in, and manage profile data.",
    )
    add_h(doc, "What it does today", 2)
    add_p(doc, "Home: showcase content for attractions, stores, and events.")
    add_p(doc, "Tickets: visual placeholder, real purchase flow still in progress.")
    add_p(doc, "Map: zoom/pan map with tappable points of interest.")
    add_p(doc, "Profile: login, registration, profile editing, photo and address updates.")
    add_h(doc, "Languages and technologies (simple terms)", 2)
    add_p(doc, "Mobile app: React Native + Expo.")
    add_p(doc, "API server: Node.js + Express.")
    add_p(doc, "Database: Microsoft SQL Server.")

    add_h(doc, "Part 2 — Features by screen", 1)
    add_h(doc, "Home", 2)
    add_p(doc, "Brand-led entry screen with carousels and shortcuts that can open filtered map views.")
    add_h(doc, "Tickets", 2)
    add_p(doc, "Placeholder screen with themed visual identity; checkout flow not enabled in this build.")
    add_h(doc, "Map", 2)
    add_p(doc, "Custom pan/zoom interactions, hotspot details, API fetch with local fallback data.")
    add_h(doc, "Profile", 2)
    add_p(doc, "Authentication flows, personal data management, and extra capabilities for employee-tagged accounts.")

    add_h(doc, "Part 3 — Architecture, engineering, and area status", 1)
    add_h(doc, "Overall architecture", 2)
    add_p(
        doc,
        "Client-server architecture: React Native app consumes a Node/Express REST API; data is stored in SQL Server.",
    )
    add_h(doc, "Current state by area", 2)
    add_p(
        doc,
        "Backend: closer to a professional structure with routes/, services/, utils/, schema/, workers/, and db/connection.js; "
        "however server.js is still larger than ideal.",
    )
    add_p(
        doc,
        "React Native app: partial extraction to src/ happened (config, types, http services, utils, map, MapaTab), "
        "but App.tsx still centralizes many responsibilities.",
    )
    add_p(
        doc,
        "Reliability/maintenance: improved with health checks, timeouts, SQL troubleshooting guidance, and .env separation. "
        "Next natural step is screen-based and hook-based modularization.",
    )
    add_h(doc, "Language types used in the project", 2)
    add_p(doc, "TypeScript and JavaScript across mobile and backend layers.")
    add_p(doc, "SQL for schema, migrations, and seed data.")
    add_p(doc, "PowerShell/BAT scripts for SQL environment setup and operations.")

    add_h(doc, "Part 4 — Honest assessment: professional level?", 1)
    add_p(
        doc,
        "Partly yes; as a whole, not yet at the level of a fully screen-modular app. "
        "It is still a serious, viable project for portfolio use and incremental production hardening.",
    )

    add_h(doc, "Straight talk (what still weighs on the codebase)", 2)
    add_p(
        doc,
        "1) Frontend: monolithic App.tsx. Impact: harder reviews and refactors. Direction: split into screens/hooks/shared components.",
    )
    add_p(
        doc,
        "2) Backend: large server.js. Impact: continued growth in one file. Direction: domain-based route modules.",
    )
    add_p(
        doc,
        "3) Limited visible automated tests. Impact: higher regression risk. Direction: API tests for critical endpoints + smoke tests.",
    )
    add_p(
        doc,
        "4) Environment complexity. Impact: onboarding mistakes. Direction: one setup guide with required variables.",
    )
    add_p(
        doc,
        "5) Many SQL scripts over time. Impact: wrong-script execution risk. Direction: standardized bootstrap flow and script headers.",
    )
    add_p(
        doc,
        "6) Observability: useful logs but not strongly structured. Direction: standardized logging and centralized error handling.",
    )
    add_p(
        doc,
        "7) Operational security: keep .env out of Git, review secrets regularly, enforce HTTPS for production.",
    )

    add_h(doc, "Final summary (EN)", 2)
    add_p(doc, "Architecture: backend is currently more modular than the app layer.")
    add_p(doc, "Maintenance: primary debt remains App.tsx and server.js size.")
    add_p(doc, "Professionalism: solid base for portfolio and incremental production evolution.")

    return doc


def _save_doc(doc: Document, root: Path, name: str) -> None:
    path = root / name
    try:
        doc.save(str(path))
        print("OK:", path)
    except OSError:
        alt = root / name.replace(".docx", "_novo.docx")
        doc.save(str(alt))
        print("OK (arquivo aberto no Word; salvo como):", alt)


def main() -> None:
    root = Path(__file__).resolve().parent
    _save_doc(build_pt(), root, "DOCFUNCIONALIDADES_PT-BR.docx")
    _save_doc(build_en(), root, "DOCFUNCIONALIDADES_EN.docx")


if __name__ == "__main__":
    main()
